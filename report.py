"""
=============================================================================
report.py  --  IN SILICO SYNTHESIS REPORT
=============================================================================

PURPOSE

A ranking table is not a record of what was done. For a design to be
reproducible - and for a reviewer to be able to check it - the following must
be recorded together:

  - the input: target, host, reference data, their versions
  - every parameter that influenced the outcome, including those left at
    their defaults
  - the metrics of each selected candidate
  - the physical forms to be ordered
  - the limitations of the method, stated rather than implied

This module assembles all of that into a single document. Two formats are
produced: Markdown for reading and version control, and HTML for printing.

The report is deliberately verbose about parameters. A design that cannot be
reproduced from its report is not a result.

Author: Antonina Jarecka
=============================================================================
"""

from dataclasses import dataclass, field
from typing import Dict, List, Optional
from datetime import datetime
import hashlib


def _hash_sequence(seq: str, n: int = 12) -> str:
    """Short SHA-256 digest - identifies the exact input without storing it."""
    return hashlib.sha256(seq.encode()).hexdigest()[:n]


@dataclass
class ReportContext:
    """Everything needed to reconstruct the run."""
    target_name: str
    target_sequence: str
    host_profile: str
    lengths: tuple
    parameters: Dict[str, object] = field(default_factory=dict)
    transcriptome_file: Optional[str] = None
    transcriptome_size_mb: Optional[float] = None
    isolate_file: Optional[str] = None
    n_isolates: Optional[int] = None
    conservation_threshold: Optional[float] = None
    conservation_method: Optional[str] = None
    tool_versions: Dict[str, str] = field(default_factory=dict)
    timestamp: str = field(
        default_factory=lambda: datetime.now().strftime('%Y-%m-%d %H:%M'))


def build_markdown(ctx: ReportContext,
                   candidates: List[Dict],
                   sensitivity: Optional[Dict] = None,
                   order_rows: Optional[List[Dict]] = None,
                   n_generated: Optional[int] = None,
                   n_filtered: Optional[Dict[str, int]] = None) -> str:
    """Assembles the full report as Markdown."""
    L: List[str] = []

    L.append('# In silico siRNA design report')
    L.append('')
    L.append(f'Generated {ctx.timestamp}')
    L.append('')
    L.append('---')
    L.append('')

    # ---------- 1. INPUT ----------
    L.append('## 1. Input')
    L.append('')
    L.append('| item | value |')
    L.append('|---|---|')
    L.append(f'| Target | {ctx.target_name} |')
    L.append(f'| Target length | {len(ctx.target_sequence)} nt |')
    L.append(f'| Target checksum (SHA-256, 12 char) | '
             f'`{_hash_sequence(ctx.target_sequence)}` |')
    L.append(f'| Host profile | {ctx.host_profile} |')
    L.append(f'| siRNA lengths | {", ".join(str(x) for x in ctx.lengths)} nt |')
    if ctx.transcriptome_file:
        L.append(f'| Off-target reference | {ctx.transcriptome_file} '
                 f'({ctx.transcriptome_size_mb} MB) |')
    else:
        L.append('| Off-target reference | **none - analysis skipped** |')
    if ctx.isolate_file:
        L.append(f'| Isolate set | {ctx.isolate_file} '
                 f'({ctx.n_isolates} sequences) |')
    L.append('')
    L.append('The checksum identifies the exact input sequence. A design '
             'cannot be reproduced without it, since a different transcript '
             'variant of the same gene gives different candidate positions.')
    L.append('')

    # ---------- 2. PARAMETERS ----------
    L.append('## 2. Parameters')
    L.append('')
    L.append('All values are listed, including those left at their defaults, '
             'so that the run can be reproduced exactly.')
    L.append('')
    L.append('| parameter | value |')
    L.append('|---|---|')
    for k, v in sorted(ctx.parameters.items()):
        L.append(f'| {k} | {v} |')
    L.append('')

    if ctx.conservation_threshold is not None:
        L.append('### Conservation threshold')
        L.append('')
        L.append(f'Value: **{ctx.conservation_threshold}**')
        L.append(f'Method: {ctx.conservation_method}')
        L.append('')

    if ctx.tool_versions:
        L.append('### Software')
        L.append('')
        for k, v in sorted(ctx.tool_versions.items()):
            L.append(f'- {k} {v}')
        L.append('')

    # ---------- 3. FILTERING ----------
    if n_generated is not None:
        L.append('## 3. Filtering')
        L.append('')
        L.append('| stage | candidates |')
        L.append('|---|---|')
        L.append(f'| Windows generated | {n_generated} |')
        if n_filtered:
            for k, v in n_filtered.items():
                L.append(f'| Removed by {k} | {v} |')
        L.append(f'| Entering scoring | {len(candidates)} |')
        L.append('')

    # ---------- 4. CANDIDATES ----------
    L.append('## 4. Selected candidates')
    L.append('')
    if not candidates:
        L.append('No candidate passed the filters.')
        L.append('')
    else:
        for c in candidates:
            L.append(f'### {c.get("nazwa", "?")}')
            L.append('')
            L.append(f'```')
            L.append(f'guide     5\'-{c.get("guide_rna", "")}-3\'')
            L.append(f'passenger 5\'-{c.get("passenger_rna", "")}-3\'')
            L.append(f'```')
            L.append('')
            L.append('| metric | value |')
            L.append('|---|---|')
            pola = [
                ('Position in target', f'{c.get("poz_start_1based")}'
                                       f'-{c.get("poz_end_1based")}'),
                ('Length', f'{c.get("dlugosc")} nt'),
                ('DCL class', c.get('klasa_DCL', 'n/a')),
                ('Seed (2-8)', c.get('seed_2_8')),
                ('GC content', f'{c.get("gc_proc")} %'),
                ('Asymmetry', f'{c.get("asymetria")} kcal/mol'),
                ('Duplex dG', f'{c.get("dG_dupleksu")} kcal/mol'),
                ('Tm', f'{c.get("Tm")} C'),
                ('Guide MFE', f'{c.get("mfe_guide")} kcal/mol'),
                ('Composite score', c.get('wynik')),
                ('Rank', c.get('ranga')),
                ('On Pareto front', 'yes' if c.get('pareto') else 'no'),
                ('Off-target status', c.get('offtarget_flaga', 'n/a')),
                ('Critical off-target hits', c.get('trafien_krytycznych', 'n/a')),
                ('Seed enrichment vs chance', c.get('wzbogacenie_seed', 'n/a')),
            ]
            if c.get('pokrycie_izolatow') is not None:
                pola.append(('Isolate coverage',
                             f'{c["pokrycie_izolatow"]:.1%}'))
            for etykieta, wartosc in pola:
                if wartosc is not None:
                    L.append(f'| {etykieta} | {wartosc} |')
            L.append('')

    # ---------- 5. SENSITIVITY ----------
    if sensitivity:
        L.append('## 5. Sensitivity to scoring weights')
        L.append('')
        L.append(f'- Baseline top-5: {", ".join(sensitivity["top_bazowy"])}')
        L.append(f'- Stable core: '
                 f'{", ".join(sensitivity["stabilny_rdzen"]) or "empty"}')
        L.append(f'- Stability: **{sensitivity["stabilnosc"]}**')
        L.append(f'- Interpretation: {sensitivity["interpretacja"]}')
        L.append('')
        L.append('Each weight was altered by 25 per cent in turn and the '
                 'ranking recomputed. A stability close to 1.0 indicates the '
                 'ordering does not depend on the arbitrary choice of '
                 'weights.')
        L.append('')

    # ---------- 6. ORDER SHEET ----------
    if order_rows:
        L.append('## 6. Sequences to order')
        L.append('')
        L.append('| candidate | form | chemistry | length | sequence 5\'-3\' |')
        L.append('|---|---|---|---|---|')
        for r in order_rows:
            L.append(f'| {r["candidate"]} | {r["form"]} | {r["chemistry"]} '
                     f'| {r["length_nt"]} | `{r["sequence_5_3"]}` |')
        L.append('')
        L.append('**Controls.** The C911 design - positions 9 to 11 replaced '
                 'by their complement - is the recommended specificity '
                 'control. It retains the seed region and therefore any '
                 'off-target activity, while the central mismatches abolish '
                 'cleavage of the intended target. Scrambled controls were '
                 'shown by Buehler et al. (2012) to lose activity in both '
                 'true- and false-positive groups and cannot distinguish '
                 'between them.')
        L.append('')

    # ---------- 7. LIMITATIONS ----------
    L.append('## 7. Limitations of this analysis')
    L.append('')
    ogr = []
    if not ctx.transcriptome_file:
        ogr.append('**No off-target analysis was performed.** No host '
                   'transcriptome was supplied, so the ranking does not '
                   'account for sequence safety. Results should not be '
                   'used for ordering without this step.')
    if not ctx.isolate_file:
        ogr.append('No conservation analysis was performed. For a variable '
                   'pathogen this means the design may not act on all '
                   'isolates.')
    ogr.extend([
        'Predictions are computational. No experimental validation of '
        'silencing efficiency was carried out.',
        'Scoring weights were set manually from the literature rather than '
        'learned from measured efficacy data.',
        'Melting temperatures carry no ionic-strength correction; the '
        'nearest-neighbour parameters apply to 1 M NaCl.',
        'Target-site accessibility is estimated from window free energy '
        'rather than from unpaired probability.',
        'Conservation analysis detects substitutions only; insertions and '
        'deletions would require a multiple sequence alignment.',
    ])
    for o in ogr:
        L.append(f'- {o}')
    L.append('')

    # ---------- 8. REFERENCES ----------
    L.append('## 8. Key references')
    L.append('')
    L.append('1. Xia T. et al. (1998) *Biochemistry* 37:14719-14735 — '
             'nearest-neighbour thermodynamic parameters')
    L.append('2. Lorenz R. et al. (2011) *Algorithms Mol Biol* 6:26 — '
             'ViennaRNA')
    L.append('3. Khvorova A. et al. (2003) *Cell* 115:209-216 — strand bias')
    L.append('4. Schwarz D.S. et al. (2003) *Cell* 115:199-208 — asymmetry')
    L.append('5. Bartel D.P. (2009) *Cell* 136:215-233 — seed classification')
    L.append('6. Buehler E. et al. (2012) *PLoS ONE* 7(12):e51942 — '
             'C911 control')
    L.append('')
    L.append('A complete reference list is given in the Methods tab of the '
             'application.')
    L.append('')
    L.append('---')
    L.append('')
    L.append('*Generated by siRNA Design Studio. Bibliographic details '
             'should be verified before citation.*')

    return '\n'.join(L)


def build_html(markdown_text: str, title: str = 'siRNA design report') -> str:
    """
    Wraps the report in printable HTML.

    A minimal Markdown conversion is used rather than an external library,
    to keep the module free of dependencies. Headings, tables, code blocks,
    lists and bold text are supported - which is all the report uses.
    """
    css = """
    @page { size: A4; margin: 20mm; }
    body { font-family: 'Georgia', 'Times New Roman', serif; font-size: 10.5pt;
           line-height: 1.55; color: #23202B; max-width: 190mm; margin: 0 auto;
           padding: 12mm; }
    h1 { font-size: 19pt; color: #4E3670; border-bottom: 2.5px solid #B8963F;
         padding-bottom: 5px; margin-bottom: 4px; }
    h2 { font-size: 14pt; color: #6B4E96; margin-top: 22px;
         border-bottom: 1px solid #D8D3E0; padding-bottom: 3px; }
    h3 { font-size: 11.5pt; color: #6B4E96; margin-top: 15px; }
    table { border-collapse: collapse; width: 100%; margin: 9px 0;
            font-size: 9pt; }
    th { background: #6B4E96; color: #fff; padding: 5px 8px; text-align: left;
         font-family: 'Helvetica', sans-serif; font-weight: 600; }
    td { padding: 4px 8px; border-bottom: 1px solid #E4E1EA; }
    tr:nth-child(even) td { background: #FAF8FC; }
    code, pre { font-family: 'Consolas', 'Courier New', monospace;
                font-size: 9pt; background: #F5F2F8; border: 1px solid #E0DBE8;
                border-radius: 3px; }
    code { padding: 1px 4px; }
    pre { padding: 8px 11px; overflow-x: auto; line-height: 1.45; }
    ul { margin: 7px 0; padding-left: 22px; }
    li { margin: 3px 0; }
    hr { border: none; border-top: 1px solid #D8D3E0; margin: 18px 0; }
    strong { color: #4E3670; }
    """

    linie = markdown_text.split('\n')
    out, w_tabeli, w_kodzie, w_liscie = [], False, False, False

    def zamknij_tabele():
        nonlocal w_tabeli
        if w_tabeli:
            out.append('</table>')
            w_tabeli = False

    def zamknij_liste():
        nonlocal w_liscie
        if w_liscie:
            out.append('</ul>')
            w_liscie = False

    def inline(t: str) -> str:
        import re
        t = (t.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'))
        t = re.sub(r'`([^`]+)`', r'<code>\1</code>', t)
        t = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', t)
        t = re.sub(r'\*([^*]+)\*', r'<em>\1</em>', t)
        return t

    for l in linie:
        if l.strip().startswith('```'):
            zamknij_tabele(); zamknij_liste()
            out.append('</pre>' if w_kodzie else '<pre>')
            w_kodzie = not w_kodzie
            continue
        if w_kodzie:
            out.append(l.replace('&', '&amp;').replace('<', '&lt;'))
            continue

        s = l.rstrip()
        if not s:
            zamknij_tabele(); zamknij_liste()
            continue
        if s.startswith('|'):
            komorki = [c.strip() for c in s.strip('|').split('|')]
            if all(set(c) <= set('-: ') for c in komorki):
                continue
            if not w_tabeli:
                out.append('<table>'); w_tabeli = True
                out.append('<tr>' + ''.join(f'<th>{inline(c)}</th>'
                                            for c in komorki) + '</tr>')
            else:
                out.append('<tr>' + ''.join(f'<td>{inline(c)}</td>'
                                            for c in komorki) + '</tr>')
            continue
        zamknij_tabele()
        if s.startswith('### '):
            zamknij_liste(); out.append(f'<h3>{inline(s[4:])}</h3>')
        elif s.startswith('## '):
            zamknij_liste(); out.append(f'<h2>{inline(s[3:])}</h2>')
        elif s.startswith('# '):
            zamknij_liste(); out.append(f'<h1>{inline(s[2:])}</h1>')
        elif s.startswith('---'):
            zamknij_liste(); out.append('<hr>')
        elif s.startswith('- ') or s.startswith('* '):
            if not w_liscie:
                out.append('<ul>'); w_liscie = True
            out.append(f'<li>{inline(s[2:])}</li>')
        else:
            zamknij_liste(); out.append(f'<p>{inline(s)}</p>')

    zamknij_tabele(); zamknij_liste()

    return (f'<!DOCTYPE html><html><head><meta charset="utf-8">'
            f'<title>{title}</title><style>{css}</style></head>'
            f'<body>{"".join(out)}</body></html>')


if __name__ == '__main__':
    ctx = ReportContext(
        target_name='GFP_ORF', target_sequence='ATGGTGAGCAAGGGC' * 48,
        host_profile='plant', lengths=(21, 22, 24),
        parameters={'gc_min': 30.0, 'gc_max': 52.0, 'min_asymmetry': 0.3,
                    'max_guide_mfe': -3.0, 'exclude_5prime_nt': 75,
                    'exclude_3prime_nt': 50},
        tool_versions={'ViennaRNA': '2.7.2', 'Python': '3.11'})

    kandydaci = [{
        'nazwa': '21nt_pos480', 'guide_rna': 'UUGAAGUUCACCUUGAUGCCG',
        'passenger_rna': 'CGGCAUCAAGGUGAACUUCAA', 'poz_start_1based': 480,
        'poz_end_1based': 500, 'dlugosc': 21, 'seed_2_8': 'UGAAGUU',
        'gc_proc': 47.6, 'asymetria': 4.83, 'dG_dupleksu': -36.71,
        'Tm': 84.9, 'mfe_guide': 0.0, 'wynik': 0.83, 'ranga': 1,
        'pareto': True, 'offtarget_flaga': 'OK', 'trafien_krytycznych': 0,
        'klasa_DCL': 'DCL4 - PTGS'}]

    md = build_markdown(ctx, kandydaci, n_generated=575,
                        n_filtered={'physicochemical filters': 560})
    print(md[:1600])
    print('...')
    html = build_html(md)
    print(f'\nHTML: {len(html)} characters')


# ============================================================================
# DOCX EXPORT
# ============================================================================

def build_docx(ctx: 'ReportContext',
               candidates: List[Dict],
               sensitivity: Optional[Dict] = None,
               order_rows: Optional[List[Dict]] = None,
               n_filtered: Optional[Dict[str, int]] = None):
    """
    Builds the report as a Word document.

    Returns the file contents as bytes, or None if python-docx is not
    installed.

    HTML printed from a browser (Ctrl+P, save as PDF) is an equally good
    route to a printable file and needs nothing installed; DOCX is offered
    because reports are often edited before circulation.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        return None

    import io as _io

    FIOLET = RGBColor(0x4E, 0x36, 0x70)
    FIOLET_J = RGBColor(0x6B, 0x4E, 0x96)

    doc = Document()

    # page setup
    for sec in doc.sections:
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)

    # base style
    st_normal = doc.styles['Normal']
    st_normal.font.name = 'Calibri'
    st_normal.font.size = Pt(10)

    def naglowek(tekst, poziom=1):
        h = doc.add_heading(tekst, level=poziom)
        for r in h.runs:
            r.font.color.rgb = FIOLET if poziom == 0 else FIOLET_J
        return h

    def akapit(tekst, rozmiar=10, kursywa=False, bold=False):
        p = doc.add_paragraph()
        r = p.add_run(tekst)
        r.font.size = Pt(rozmiar)
        r.italic = kursywa
        r.bold = bold
        return p

    def mono(tekst):
        p = doc.add_paragraph()
        r = p.add_run(tekst)
        r.font.name = 'Consolas'
        r.font.size = Pt(9)
        return p

    def tabela(naglowki, wiersze):
        t = doc.add_table(rows=1, cols=len(naglowki))
        t.style = 'Light Grid Accent 1'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(naglowki):
            kom = t.rows[0].cells[i]
            kom.text = ''
            r = kom.paragraphs[0].add_run(str(h))
            r.bold = True
            r.font.size = Pt(9)
        for w in wiersze:
            komorki = t.add_row().cells
            for i, v in enumerate(w):
                komorki[i].text = ''
                r = komorki[i].paragraphs[0].add_run(str(v))
                r.font.size = Pt(8.5)
                if isinstance(v, str) and len(v) > 15 and set(v) <= set('ACGUTdN'):
                    r.font.name = 'Consolas'
        return t

    # ---------- title ----------
    naglowek('In silico siRNA design report', 0)
    akapit(f'Generated {ctx.timestamp}', 9, kursywa=True)

    # ---------- 1 input ----------
    naglowek('1. Input', 1)
    wiersze = [
        ['Target', ctx.target_name],
        ['Target length', f'{len(ctx.target_sequence)} nt'],
        ['Checksum (SHA-256, 12 char)', _hash_sequence(ctx.target_sequence)],
        ['Host profile', ctx.host_profile],
        ['siRNA lengths', ', '.join(str(x) for x in ctx.lengths) + ' nt'],
        ['Off-target reference',
         f'{ctx.transcriptome_file} ({ctx.transcriptome_size_mb} MB)'
         if ctx.transcriptome_file else 'none - analysis skipped'],
    ]
    if ctx.isolate_file:
        wiersze.append(['Isolate set',
                        f'{ctx.isolate_file} ({ctx.n_isolates} sequences)'])
    tabela(['Item', 'Value'], wiersze)
    akapit('The checksum identifies the exact input sequence. A design cannot '
           'be reproduced without it, since a different transcript variant of '
           'the same gene gives different candidate positions.', 9,
           kursywa=True)

    # ---------- 2 parameters ----------
    naglowek('2. Parameters', 1)
    akapit('All values are listed, including those left at their defaults, so '
           'that the run can be reproduced exactly.', 9, kursywa=True)
    tabela(['Parameter', 'Value'],
           [[k, str(v)] for k, v in sorted(ctx.parameters.items())])

    if ctx.conservation_threshold is not None:
        naglowek('Conservation threshold', 2)
        akapit(f'Value: {ctx.conservation_threshold}   ·   '
               f'Method: {ctx.conservation_method}')

    # ---------- 3 candidates ----------
    naglowek('3. Selected candidates', 1)
    if not candidates:
        akapit('No candidate passed the filters.')
    for c in candidates:
        naglowek(str(c.get('nazwa', '?')), 2)
        mono(f"guide     5'-{c.get('guide_rna', '')}-3'")
        mono(f"passenger 5'-{c.get('passenger_rna', '')}-3'")
        pola = [
            ['Position', f"{c.get('poz_start_1based')}-{c.get('poz_end_1based')}"],
            ['Length', f"{c.get('dlugosc')} nt"],
            ['DCL class', c.get('klasa_DCL', 'n/a')],
            ['Seed (2-8)', c.get('seed_2_8')],
            ['GC content', f"{c.get('gc_proc')} %"],
            ['Asymmetry', f"{c.get('asymetria')} kcal/mol"],
            ['Duplex dG', f"{c.get('dG_dupleksu')} kcal/mol"],
            ['Tm', f"{c.get('Tm')} C"],
            ['Composite score', c.get('wynik')],
            ['Rank', c.get('ranga')],
            ['Pareto front', 'yes' if c.get('pareto') else 'no'],
            ['Off-target status', c.get('offtarget_flaga', 'n/a')],
            ['Critical off-target hits', c.get('trafien_krytycznych', 'n/a')],
        ]
        if c.get('pokrycie_izolatow') is not None:
            pola.append(['Isolate coverage', f"{c['pokrycie_izolatow']:.1%}"])
        tabela(['Metric', 'Value'],
               [[a, str(b)] for a, b in pola if b is not None])

    # ---------- 4 sensitivity ----------
    if sensitivity:
        naglowek('4. Sensitivity to scoring weights', 1)
        tabela(['Item', 'Value'], [
            ['Baseline top-5', ', '.join(sensitivity['top_bazowy'])],
            ['Stable core', ', '.join(sensitivity['stabilny_rdzen']) or 'empty'],
            ['Stability', str(sensitivity['stabilnosc'])],
            ['Interpretation', sensitivity['interpretacja']]])
        akapit('Each weight was altered by 25 per cent in turn and the ranking '
               'recomputed. A stability close to 1.0 indicates the ordering '
               'does not depend on the arbitrary choice of weights.', 9,
               kursywa=True)

    # ---------- 5 order sheet ----------
    if order_rows:
        naglowek('5. Sequences to order', 1)
        tabela(['Candidate', 'Form', 'Chem.', 'nt', "Sequence 5'-3'"],
               [[r['candidate'], r['form'], r['chemistry'],
                 r['length_nt'], r['sequence_5_3']] for r in order_rows])
        akapit('Controls. The C911 design - positions 9 to 11 replaced by '
               'their complement - is the recommended specificity control. It '
               'retains the seed region and therefore any off-target activity, '
               'while the central mismatches abolish cleavage of the intended '
               'target. Scrambled controls were shown by Buehler et al. (2012) '
               'to lose activity in both true- and false-positive groups and '
               'cannot distinguish between them.', 9)

    # ---------- 6 limitations ----------
    naglowek('6. Limitations of this analysis', 1)
    ogr = []
    if not ctx.transcriptome_file:
        ogr.append('No off-target analysis was performed. No host '
                   'transcriptome was supplied, so the ranking does not '
                   'account for sequence safety. Results should not be used '
                   'for ordering without this step.')
    if not ctx.isolate_file:
        ogr.append('No conservation analysis was performed. For a variable '
                   'pathogen the design may not act on all isolates.')
    ogr += [
        'Predictions are computational. No experimental validation of '
        'silencing efficiency was carried out.',
        'Scoring weights were set manually from the literature rather than '
        'learned from measured efficacy data.',
        'Melting temperatures carry no ionic-strength correction.',
        'Target-site accessibility is estimated from window free energy.',
        'Conservation analysis detects substitutions only.',
    ]
    for o in ogr:
        doc.add_paragraph(o, style='List Bullet')

    # ---------- 7 references ----------
    naglowek('7. Key references', 1)
    for r in [
        'Xia T. et al. (1998) Biochemistry 37:14719-14735 - nearest-neighbour '
        'thermodynamic parameters',
        'Lorenz R. et al. (2011) Algorithms Mol Biol 6:26 - ViennaRNA',
        'Khvorova A. et al. (2003) Cell 115:209-216 - strand bias',
        'Schwarz D.S. et al. (2003) Cell 115:199-208 - asymmetry',
        'Bartel D.P. (2009) Cell 136:215-233 - seed classification',
        'Buehler E. et al. (2012) PLoS ONE 7(12):e51942 - C911 control',
    ]:
        doc.add_paragraph(r, style='List Number')

    akapit('Generated by siRNA Design Studio. Bibliographic details should be '
           'verified before citation.', 8, kursywa=True)

    bufor = _io.BytesIO()
    doc.save(bufor)
    return bufor.getvalue()
